import io
from datetime import datetime, date, timedelta
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from app.models.report import Report
from app.models.task import Task
from app.models.email import Email
from app.models.work_log import WorkLog
from app.models.user import User


async def _collect_weekly_data(db: AsyncSession) -> dict:
    today = date.today()
    week_start = today - timedelta(days=today.weekday())
    week_end = today

    done_tasks = await db.scalar(select(func.count(Task.id)).where(
        Task.status == "done",
        func.date(Task.done_at) >= week_start,
        func.date(Task.done_at) <= week_end
    ))
    overdue = await db.scalar(select(func.count(Task.id)).where(
        Task.status != "done", Task.due_date < today
    ))
    emails_done = await db.scalar(select(func.count(Email.id)).where(Email.status == "done"))

    logs_result = await db.execute(select(WorkLog).where(
        WorkLog.log_date >= week_start, WorkLog.log_date <= week_end
    ))
    logs = logs_result.scalars().all()

    issues = [l.issues for l in logs if l.issues]
    contents = [l.content for l in logs if l.content]
    next_plans = [l.next_plan for l in logs if getattr(l, "next_plan", None)]

    return {
        "period_start": str(week_start),
        "period_end": str(week_end),
        "done_tasks": done_tasks or 0,
        "overdue_tasks": overdue or 0,
        "emails_processed": emails_done or 0,
        "completed_work": contents,
        "issues": issues,
        "next_plans": next_plans,
        "generated_at": str(datetime.now()),
    }


async def _collect_monthly_data(db: AsyncSession) -> dict:
    today = date.today()
    month_start = today.replace(day=1)

    users_result = await db.execute(select(User).where(User.is_active == True))
    users = users_result.scalars().all()

    done_result = await db.execute(
        select(Task.assigned_to_id, func.count(Task.id)).where(
            Task.status == "done", func.date(Task.done_at) >= month_start
        ).group_by(Task.assigned_to_id)
    )
    done_stats = {row[0]: row[1] for row in done_result}

    in_progress_result = await db.execute(
        select(Task.assigned_to_id, func.count(Task.id)).where(
            Task.status.in_(["todo", "in_progress", "review"])
        ).group_by(Task.assigned_to_id)
    )
    in_progress_stats = {row[0]: row[1] for row in in_progress_result}

    user_stats = []
    for u in users:
        user_stats.append({"name": u.name, "done": done_stats.get(u.id, 0), "in_progress": in_progress_stats.get(u.id, 0)})

    total_done = await db.scalar(select(func.count(Task.id)).where(
        Task.status == "done", func.date(Task.done_at) >= month_start
    ))
    total_tasks = await db.scalar(select(func.count(Task.id)).where(
        func.date(Task.created_at) >= month_start
    ))

    return {
        "period": str(today.strftime("%Y-%m")),
        "total_done_tasks": total_done or 0,
        "total_tasks": total_tasks or 0,
        "deadline_rate": round((total_done / total_tasks * 100) if total_tasks else 0),
        "user_stats": user_stats,
        "generated_at": str(datetime.now()),
    }


async def generate_weekly(db: AsyncSession, user_id: int):
    today = date.today()
    period = f"{today.isocalendar()[0]}-W{today.isocalendar()[1]:02d}"
    data = await _collect_weekly_data(db)
    return await _save_report(db, "weekly", period, data, user_id)


async def generate_monthly(db: AsyncSession, user_id: int):
    today = date.today()
    period = today.strftime("%Y-%m")
    data = await _collect_monthly_data(db)
    return await _save_report(db, "monthly", period, data, user_id)


async def _save_report(db: AsyncSession, type_: str, period: str, content: dict, user_id: int):
    result = await db.execute(select(Report).where(Report.type == type_, Report.period == period))
    report = result.scalar_one_or_none()
    if report:
        report.content = content
    else:
        report = Report(type=type_, period=period, content=content, created_by_id=user_id)
        db.add(report)
    await db.commit()
    await db.refresh(report)
    return _r(report)


async def list_reports(db: AsyncSession, type_: str | None):
    q = select(Report)
    if type_:
        q = q.where(Report.type == type_)
    q = q.order_by(Report.generated_at.desc())
    result = await db.execute(q)
    return [_r(r) for r in result.scalars().all()]


async def get_report(db: AsyncSession, report_id: int):
    result = await db.execute(select(Report).where(Report.id == report_id))
    r = result.scalar_one_or_none()
    return _r(r) if r else None


async def update_report(db: AsyncSession, report_id: int, content: dict):
    result = await db.execute(select(Report).where(Report.id == report_id))
    r = result.scalar_one_or_none()
    if r:
        r.content = content
        await db.commit()
        await db.refresh(r)
    return _r(r)


async def export_to_docx(db: AsyncSession, report_id: int) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor
    r = await get_report(db, report_id)
    doc = Document()
    doc.add_heading(f"{'주간' if r['type'] == 'weekly' else '월간'}업무보고 — {r['period']}", 0)
    content = r["content"]

    if r["type"] == "weekly":
        doc.add_heading("실적 요약", 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "완료 태스크", "지연 태스크", "처리 이메일"
        row = table.add_row().cells
        row[0].text = str(content.get("done_tasks", 0))
        row[1].text = str(content.get("overdue_tasks", 0))
        row[2].text = str(content.get("emails_processed", 0))

        doc.add_heading("완료 업무", 1)
        for item in content.get("completed_work", []):
            doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("이슈 / 리스크", 1)
        for item in content.get("issues", []):
            doc.add_paragraph(item, style="List Bullet")

        doc.add_heading("다음 업무 계획", 1)
        for item in content.get("next_plans", []):
            doc.add_paragraph(item, style="List Bullet")

    elif r["type"] == "monthly":
        doc.add_heading("팀원별 업무 현황", 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "이름", "완료", "진행중"
        for us in content.get("user_stats", []):
            row = table.add_row().cells
            row[0].text = us["name"]
            row[1].text = str(us["done"])
            row[2].text = str(us["in_progress"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


async def export_to_pdf(db: AsyncSession, report_id: int) -> io.BytesIO:
    from weasyprint import HTML
    r = await get_report(db, report_id)
    c = r["content"]
    title = f"{'주간' if r['type'] == 'weekly' else '월간'}업무보고 — {r['period']}"

    def esc(s):
        return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    sections = ""
    if r["type"] == "weekly":
        sections += f"""
        <h2>실적 요약</h2>
        <table>
          <tr><th>완료 태스크</th><th>지연 태스크</th><th>처리 이메일</th></tr>
          <tr><td>{c.get('done_tasks',0)}</td><td>{c.get('overdue_tasks',0)}</td><td>{c.get('emails_processed',0)}</td></tr>
        </table>"""
        for sec_title, key in [("완료 업무", "completed_work"), ("이슈 / 리스크", "issues"), ("다음 업무 계획", "next_plans")]:
            items = c.get(key, []) or []
            lis = "".join(f"<li>{esc(i)}</li>" for i in items) or "<li class='empty'>없음</li>"
            sections += f"<h2>{sec_title}</h2><ul>{lis}</ul>"
    else:
        rows = "".join(
            f"<tr><td>{esc(u['name'])}</td><td>{u['done']}</td><td>{u['in_progress']}</td></tr>"
            for u in c.get("user_stats", [])
        )
        sections += f"""
        <h2>월간 요약</h2>
        <table>
          <tr><th>완료 태스크</th><th>전체 태스크</th><th>마감 준수율</th></tr>
          <tr><td>{c.get('total_done_tasks',0)}</td><td>{c.get('total_tasks',0)}</td><td>{c.get('deadline_rate',0)}%</td></tr>
        </table>
        <h2>팀원별 현황</h2>
        <table><tr><th>이름</th><th>완료</th><th>진행중</th></tr>{rows}</table>"""

    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>
      body {{ font-family: 'Noto Sans CJK KR', 'Malgun Gothic', sans-serif; color: #1e293b; padding: 10px; }}
      h1 {{ font-size: 22px; border-bottom: 2px solid #334155; padding-bottom: 8px; }}
      h2 {{ font-size: 15px; margin-top: 22px; color: #334155; }}
      table {{ border-collapse: collapse; width: 100%; margin-top: 8px; }}
      th, td {{ border: 1px solid #cbd5e1; padding: 6px 10px; text-align: left; font-size: 12px; }}
      th {{ background: #f1f5f9; }}
      ul {{ padding-left: 18px; }} li {{ font-size: 12px; margin: 3px 0; }}
      .empty {{ color: #94a3b8; }}
      .meta {{ color: #94a3b8; font-size: 11px; margin-top: 20px; }}
    </style></head><body>
      <h1>{title}</h1>
      {sections}
      <div class="meta">생성: {esc(c.get('generated_at',''))}</div>
    </body></html>"""

    buf = io.BytesIO()
    HTML(string=html).write_pdf(buf)
    buf.seek(0)
    return buf


def _r(r: Report) -> dict:
    if not r:
        return None
    return {"id": r.id, "type": r.type, "period": r.period,
            "content": r.content, "generated_at": str(r.generated_at)}
